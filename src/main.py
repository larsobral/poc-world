from docx import Document

def replace_university_name(doc, new_name):
    palavra_chave = 'XXXXXXXXXXXXX'
    for paragraph in doc.paragraphs:
        if palavra_chave in paragraph.text:
            paragraph.text = paragraph.text.replace(palavra_chave, new_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if palavra_chave in cell.text:
                    cell.text = cell.text.replace(palavra_chave, new_name)

# Lista com os nomes das universidades
universities = [
    "Universidade UNFBV",
    # Adicione os outros nomes de universidades aqui
]
path_doc = r'../docs/REGULAMENTO NUCLEO DE INOVACAO E TECNOLOGIA – NIT.docx'

# Itera através das universidades e cria um novo documento para cada uma
for university_name in universities:
    doc = Document(path_doc)  # Carrega o documento dentro do loop
    replace_university_name(doc, university_name)
    output_path = f'../output/output_{university_name}.docx'
    doc.save(output_path)
    print(f"Substituição de nomes concluída para {university_name}. Salvo em {output_path}")

